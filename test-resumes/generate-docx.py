"""Generate a funny fake DOCX resume for form upload testing."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Narrow margins
for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(2)

GREEN = RGBColor(0x00, 0x86, 0x3F)
GRAY = RGBColor(0x55, 0x55, 0x55)

# Header
name = doc.add_paragraph()
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = name.add_run('Cornelius "Big Mow" McGrassington III')
run.bold = True
run.font.size = Pt(22)

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = contact.add_run("123 Dandelion Lane, Weedville, NY 14999  |  (585) 555-LAWN  |  bigmow@totallyrealmail.biz")
run.font.size = Pt(10)
run.font.color.rgb = GRAY

# Divider
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(4)
run = p.add_run("_" * 85)
run.font.color.rgb = GREEN
run.font.size = Pt(6)

def add_section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = GREEN

def add_job(title, dates, company, location, bullets):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(title)
    run.bold = True
    run = p.add_run(f"    {dates}")
    run.italic = True
    run.font.color.rgb = GRAY

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    run = p2.add_run(f"{company}")
    run.bold = True
    run = p2.add_run(f" — {location}")
    run.italic = True

    for bullet in bullets:
        bp = doc.add_paragraph(bullet, style="List Bullet")
        bp.paragraph_format.space_after = Pt(1)

# Content
add_section_heading("OBJECTIVE")
doc.add_paragraph(
    'To leverage my 47 years of grass-whispering experience in a dynamic, fast-paced '
    'mowing environment where I can finally prove to my mother that landscaping is a real career.'
)

add_section_heading("PROFESSIONAL EXPERIENCE")
add_job("Chief Grass Officer", "2019 – Present", "McGrassington Family Lawn Empire", "Backyard, NY", [
    "Personally responsible for maintaining 0.3 acres of residential turf to competition-grade standards",
    'Developed proprietary "just eyeball it" edging technique, achieving 40% straighter lines than previous seasons',
    "Reduced dandelion population by 12% through aggressive staring and motivational speeches",
    "Managed a team of one (myself) with zero HR complaints",
])
add_job("Senior Leaf Relocation Specialist", "2015 – 2019", "Dave's Definitely Legit Landscaping", "Rochester, NY", [
    "Relocated over 2.3 million leaves across 4 autumn seasons using state-of-the-art rake technology",
    'Pioneered the "blow it into the neighbor\'s yard" efficiency methodology (later discontinued)',
    "Employee of the Month, November 2017 (only employee)",
])
add_job("Junior Weed Identification Intern", "2012 – 2015", "Rochester Community College Grounds Department", "", [
    'Correctly identified grass as "green" 98.7% of the time',
    'Maintained detailed spreadsheet of which weeds looked "kinda cool actually"',
    "Assisted in the Great Mulch Incident of 2014 (details available upon request)",
])

add_section_heading("EDUCATION")
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
run = p.add_run("B.S. in Turf Sciences (Incomplete)")
run.bold = True
run = p.add_run("    2008 – 2012")
run.italic = True
run.font.color.rgb = GRAY
doc.add_paragraph("Rochester Community College")
for b in [
    'GPA: 2.1 ("but I was on the Dean\'s radar, which is almost the same as the Dean\'s List")',
    "Relevant coursework: Introduction to Dirt, Advanced Dirt, The Philosophy of Mulch",
]:
    bp = doc.add_paragraph(b, style="List Bullet")
    bp.paragraph_format.space_after = Pt(1)

add_section_heading("SKILLS & CERTIFICATIONS")
for b in [
    "Equipment: Push mower (expert), riding mower (intermediate), scissors (beginner)",
    "Software: Microsoft Excel (can open it), Google Maps (for finding lawns)",
    "Languages: English, Conversational Plant",
    "Certifications: CPR (expired 2016), NYS DEC Pesticide Applicator (pending since 2019), Forklift (why not)",
]:
    bp = doc.add_paragraph(b, style="List Bullet")
    bp.paragraph_format.space_after = Pt(1)

add_section_heading("REFERENCES")
for b in [
    'My Mom — "He\'s very talented, you should hire him" (unverified)',
    "Dave from Dave's Definitely Legit Landscaping — Currently unreachable (unrelated legal matter)",
    'Dr. Susan Greenthumb, PhD — "I have never heard of this person" (Rochester Community College)',
]:
    bp = doc.add_paragraph(b, style="List Bullet")
    bp.paragraph_format.space_after = Pt(1)

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
run = p.add_run("This resume was printed on 100% recycled grass clippings. Please do not eat.")
run.italic = True
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.save(r"C:\Users\callm\Documents\PROJECTS\Westside-Website-v3\test-resumes\cornelius-mcgrassington-resume.docx")
print("DOCX saved.")
